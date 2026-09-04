"""简历文本/结构化解析。

提供两层 API：
- extract_resume_text(bytes, filename) -> str  原文本（已存在）
- parse_resume_structure(text) -> dict 结构化结果，**所有字段都来自原文 verbatim**：
  {
    "profile": {"name": str|None, "phone": str|None, "email": str|None, "years": str|None},
    "work_experiences": [
      {"company": str, "role": str, "period": str, "bullets": [str, ...]},
      ...
    ],
  }

解析策略：纯规则（正则 + 启发式）。故意不用 LLM 是为了让"原文保真"零漂移——
LLM 提取会被规范化（"ByteDance"→"字节跳动"、"3 yrs"→"3年"），而这正是用户要避免的。

针对中文技术简历常见版式做了适配：
- 工作经历 section header: "工作经历" / "Work Experience"
- 岗位 header: 通常含时间段（"2023.07 - 至今" / "2020.03 - 2022.06"）
- 公司/岗位分隔符: 竖线 / 多个空格 / 单行多列
- bullet: "• · - *" 开头，或裸行（位于 header 之后、下一个 header 之前）
"""
import io
import re
from typing import Any

from pypdf import PdfReader
import docx


# ============================================================================
# 文本提取
# ============================================================================

def extract_text_from_pdf(content_bytes: bytes) -> str:
    pdf_file = io.BytesIO(content_bytes)
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def extract_text_from_docx(content_bytes: bytes) -> str:
    docx_file = io.BytesIO(content_bytes)
    doc = docx.Document(docx_file)
    text = ""
    for para in doc.paragraphs:
        if para.text:
            text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text]
            if row_text:
                text += " | ".join(row_text) + "\n"
    return text


def extract_resume_text(content_bytes: bytes, filename: str) -> str:
    ext = filename.split('.')[-1].lower() if '.' in filename else ""
    if ext == "pdf":
        return extract_text_from_pdf(content_bytes)
    elif ext == "docx":
        return extract_text_from_docx(content_bytes)
    else:
        raise ValueError("不支持的文件格式，仅支持 PDF 和 DOCX")


# ============================================================================
# 结构化解析（核心）
# ============================================================================

# 时间段: 2023.07 - 至今 / 2020.03 - 2022.06 / 2022.09 ~ 2023.06 / 2018-2020
# 也支持 2023年07月 这种中文写法
_DATE_RE = re.compile(
    r"((?:19|20)\d{2})"
    r"(?:\s*[\.\-/年]\s*(\d{1,2})(?:\s*月)?)?"
    r"\s*[-—–~～至]+\s*"
    r"("
    r"(?:19|20)?\d{2}"
    r"(?:\s*[\.\-/年]\s*\d{1,2}(?:\s*月)?)?"
    r"|至今|present|现在|今"
    r")",
    re.IGNORECASE,
)

# 手机号
_PHONE_RE = re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)")
# 邮箱
_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
# 工作年限：必须含"经验"或"工作/从事/从业"语义词，避免把"1995年"这种出生年误识别
_YEARS_RE = re.compile(
    r"(?:工作|从事|有|累计|拥有)?\s*"
    r"(\d{1,2})\s*年"
    r"(?:\s*(\d{1,2})\s*个?月)?"
    r"(?=\s*(?:经验|工作经验|行业经验|从业|以上|$))",
    re.IGNORECASE,
)
# 候选人姓名常见位置: 简历首行的纯中文 2-4 字
_NAME_LINE_RE = re.compile(r"^[\s　]*([一-龥]{2,4})[\s　]*$")

# Section headers
_WORK_SECTION_KEYS = (
    "工作经历", "工作经验", "职业经历", "任职经历", "工作履历", "从业经历", "工作/实习经历", "工作与实习经历",
    "employment", "work experience", "professional experience", "experience",
)
_INTERN_SECTION_KEYS = (
    "实习经历", "实习经验", "实习履历", "校内实习",
    "internship", "internships", "intern experience",
)
_EDU_SECTION_KEYS = (
    "教育背景", "教育经历", "教育信息", "学历信息", "学业背景",
    "education", "academic background", "education history",
)
_PROJ_SECTION_KEYS = (
    "项目经历", "项目经验", "个人项目", "核心项目", "项目案例", "主要项目", "项目实战", "项目背景",
    "projects", "personal projects", "key projects", "project experience",
)
_SKILL_SECTION_KEYS = (
    "专业技能", "技能清单", "核心技能", "技术栈", "技能专长", "专业能力", "it技能", "技能",
    "skills", "technical skills", "professional skills", "core skills",
)
_SUMMARY_SECTION_KEYS = (
    "自我评价", "个人评价", "个人总结", "综合评价", "个人优势", "职业目标", "关于我",
    "summary", "personal summary", "about me", "profile summary",
)
_NEXT_SECTION_KEYS = (
    "教育背景", "教育经历", "教育信息", "学历信息", "学业背景",
    "项目经历", "项目经验", "个人项目", "核心项目", "项目案例", "主要项目", "项目实战",
    "专业技能", "技能清单", "核心技能", "技术栈", "技能专长", "专业能力", "it技能", "技能",
    "自我评价", "个人评价", "个人总结", "综合评价", "个人优势", "职业目标", "关于我",
    "兴趣爱好", "获奖", "证书", "语言", "荣誉",
    "education", "projects", "skills", "awards", "certifications",
    "summary", "interests",
    "工作经历", "工作经验", "职业经历", "任职经历", "工作履历", "从业经历", "工作/实习经历",
    "实习经历", "实习经验", "internship",
)

_BULLET_PREFIX_RE = re.compile(r"^[\s　]*[•·●○■□◆◇▶▷\-–—\*]\s*")
_NUMERIC_PREFIX_RE = re.compile(r"^[\s　]*(?:\d+[\.\)、]|\(\d+\)|[①-⑩])\s*")
_PUA_PREFIX_RE = re.compile(r"^[- -⁯　-〿 - ]+")


def _clean_line(s: str) -> str:
    """去掉行首的 PUA / 装饰字符 + 全部空白（含全角空格）。"""
    s = _PUA_PREFIX_RE.sub("", s)
    return s.strip()


def _clean_header_line(s: str) -> str:
    """剥除【】、[]、()、序号（1. 1、 一、）、符号（◆ ■ #）及首尾空白。"""
    s = _PUA_PREFIX_RE.sub("", s).strip()
    s = re.sub(r"^[【\[\(\{\#\*◆■●★▲\s\d一二三四五六七八九十0-9.、\-+—–]+", "", s)
    s = re.sub(r"[】\]\)\}\s]+.*$", "", s)
    return s.strip().lower()


def _is_header_match(line: str, keys: tuple[str, ...]) -> bool:
    raw = line.strip().lower()
    if not raw:
        return False
    cleaned = _clean_header_line(line)
    for k in keys:
        k_lower = k.lower()
        if cleaned == k_lower or cleaned.startswith(k_lower):
            return True
        if raw == k_lower or raw.startswith(k_lower + " ") or raw.startswith(k_lower + ":") or raw.startswith(k_lower + "："):
            return True
    return False


def _find_section_bounds(
    lines: list[str],
    start_keys: tuple[str, ...],
    end_keys: tuple[str, ...] = _NEXT_SECTION_KEYS,
) -> tuple[int | None, int]:
    """找到 start_keys 标记的 section 起始行（含）和结束行（不含）。"""
    start: int | None = None
    for i, line in enumerate(lines):
        if _is_header_match(line, start_keys):
            start = i
            break
    if start is None:
        return None, len(lines)
    for j in range(start + 1, len(lines)):
        if _is_header_match(lines[j], end_keys):
            return start, j
    return start, len(lines)


def _parse_profile(lines: list[str], full_text: str) -> dict:
    """从全文抽出基础字段：name / phone / email / years / gender / age / location。"""
    profile: dict[str, Any] = {
        "name": None, "phone": None, "email": None, "years": None,
        "gender": None, "age": None, "location": None,
    }

    # phone / email 跨全文搜
    m = _PHONE_RE.search(full_text)
    if m:
        profile["phone"] = m.group(0)
    m = _EMAIL_RE.search(full_text)
    if m:
        profile["email"] = m.group(0)
    m = _YEARS_RE.search(full_text)
    if m:
        years, months = m.group(1), m.group(2)
        if years and months:
            profile["years"] = f"{years}年{months}个月"
        elif years:
            profile["years"] = f"{years}年"
    elif "应届" in full_text or "校招" in full_text:
        profile["years"] = "应届生"

    # 性别 / 年龄 / 城市 匹配
    m_gender = re.search(r"性别[\s:：]*([男|女])", full_text)
    if m_gender:
        profile["gender"] = m_gender.group(1).strip()

    m_age = re.search(r"年龄[\s:：]*(\d{1,2}\s*岁?)", full_text)
    if m_age:
        profile["age"] = m_age.group(1).strip()

    m_loc = re.search(r"(?:城市|现居地|所在地)[\s:：]*([一-龥]{2,6})", full_text)
    if m_loc:
        profile["location"] = m_loc.group(1).strip()

    # 1. 优先提取 "姓名：张三" / "姓名 张三" / "姓名:张三"
    m_name = re.search(r"姓名[\s:：]*([一-龥]{2,4})", full_text)
    if m_name:
        profile["name"] = m_name.group(1).strip()
    else:
        # 2. 启发式兜底：前 10 行内寻找单行 2-4 字中文
        _EXCLUDE_NAME_HEADERS = {"基本信息", "个人信息", "简历信息", "个人简历", "求职意向", "基本资料", "教育背景", "工作经历", "项目经历", "专业技能", "联系方式"}
        for line in lines[:10]:
            clean_l = _clean_line(line.strip())
            m = _NAME_LINE_RE.match(clean_l)
            if m:
                val = m.group(1).strip()
                if val not in _EXCLUDE_NAME_HEADERS:
                    profile["name"] = val
                    break

    return profile


def _parse_header_company_role(line: str, date_match) -> tuple[str, str]:
    """从 header 行里抽 (company, role)。

    常见形式：
      字节跳动 | 高级后端开发工程师 | 2023.07 - 至今
      字节跳动  高级后端开发工程师  2023.07 - 至今
      字节跳动 高级后端开发工程师 2023.07 - 至今
    抽完 date 后，按 | / 多个空格 / tab 切分；**不再做关键词切分**——
    关键词切分会把"高级后端开发工程师"从中间劈开（错把"工程师"当切点）。
    切不开时整段留 company，role 留空，由 LLM 优化 bullets 阶段补回。
    """
    if date_match:
        before = line[: date_match.start()].rstrip(" \t|-—–·")
    else:
        before = line.rstrip()

    if not before:
        return "", ""

    # 优先级 1: 竖线分隔
    for sep in ("|", "｜"):
        if sep in before:
            parts = [p.strip() for p in before.split(sep) if p.strip()]
            if len(parts) >= 2:
                return parts[0], parts[1]
            if len(parts) == 1:
                return parts[0], ""

    # 优先级 2: 2+ 空格 / tab 分隔（DOCX 多列布局常见）
    parts = re.split(r"[\s　]{2,}|\t+", before)
    parts = [p.strip() for p in parts if p.strip()]
    if len(parts) >= 2:
        return parts[0], parts[1]

    # 兜底: 整段当 company，role 留空（不强行用关键词切，避免劈开长 role）
    return before, ""


def _parse_work_experiences(lines: list[str]) -> list[dict]:
    """从工作经历 section 解析出 jobs 列表。"""
    # 找到所有包含 date 的行（每行可能多个 date，取最长的那个）
    header_line_idx: list[int] = []
    header_date: list[re.Match] = []
    for i, line in enumerate(lines):
        # 一行内可能有多段工作时间（如"2023.07 - 至今\n2020.03 - 2022.06"被压在一行）
        # 取最右或最长的那个当作本 header 的 period
        matches = list(_DATE_RE.finditer(line))
        if matches:
            # 选最长的一个（更可能是完整 period）
            best = max(matches, key=lambda m: m.end() - m.start())
            header_line_idx.append(i)
            header_date.append(best)

    if not header_line_idx:
        return []

    jobs: list[dict] = []
    for n, idx in enumerate(header_line_idx):
        line = lines[idx]
        date_m = header_date[n]
        period = date_m.group(0).strip()

        # === company / role 提取 ===
        # 1) 在同一行
        company, role = _parse_header_company_role(line, date_m)
        # 2) 不在同一行: 向上 1-2 行找（DOCX 经常"公司名\n岗位\n时间" 或 "公司名 时间\n岗位"）
        if not company and idx > 0:
            for back in range(1, 3):
                if idx - back < 0:
                    break
                prev = lines[idx - back].strip()
                if not prev or _BULLET_PREFIX_RE.match(prev) or _DATE_RE.search(prev):
                    continue
                # 上一行就是 company
                company = prev
                # 再上一行可能是 role
                if not role and idx - back - 1 >= 0:
                    up = lines[idx - back - 1].strip()
                    if up and not _BULLET_PREFIX_RE.match(up) and not _DATE_RE.search(up):
                        role = up
                break
        # 3) 向下 1 行找 role（"公司名 时间\n岗位" 格式）
        if company and not role:
            next_idx = idx + 1
            if next_idx < len(lines):
                nxt = lines[next_idx].strip()
                # 不能是下一个 header（避免误把下家的 role 算进来）
                if next_idx not in header_line_idx and nxt and not _BULLET_PREFIX_RE.match(nxt):
                    role = nxt
                    # 同时把 next_idx 从 bullets 起点里跳过
                    skip_next = next_idx
                else:
                    skip_next = None
            else:
                skip_next = None
        else:
            skip_next = None

        # === bullets 提取：header 之后、下一个 header 之前 ===
        end = header_line_idx[n + 1] if n + 1 < len(header_line_idx) else len(lines)
        
        # 1. 收集当前工作经历的所有候选行，跳过空行与下一个 section header，处理多行 header 防御
        raw_job_lines: list[str] = []
        for j in range(idx + 1, end):
            if skip_next is not None and j == skip_next:
                continue
            raw = lines[j].strip()
            if not raw:
                continue
            # 跳过 section header
            if any(raw.lower().startswith(k) for k in _NEXT_SECTION_KEYS):
                break
            # 多行 header 防御
            if not _BULLET_PREFIX_RE.match(raw):
                next_non_empty = ""
                for k in range(j + 1, min(len(lines), j + 4)):
                    nxt_k = lines[k].strip()
                    if not nxt_k:
                        continue
                    next_non_empty = nxt_k
                    break
                if next_non_empty and _DATE_RE.search(next_non_empty):
                    continue
            raw_job_lines.append(raw)

        # 2. 合并因为 PDF/Word 排版自动换行而被截断/拆碎的连续行（例如行末没有标点符号且下一行不以 bullet 前缀开头）
        merged_job_lines: list[str] = []
        for line in raw_job_lines:
            if not merged_job_lines:
                merged_job_lines.append(line)
                continue
            
            prev_line = merged_job_lines[-1]
            has_bullet_prefix = bool(_BULLET_PREFIX_RE.match(line))
            has_numeric_prefix = bool(_NUMERIC_PREFIX_RE.match(line))
            
            prev_clean = prev_line.strip()
            ends_with_sentence_terminator = False
            if prev_clean:
                last_char = prev_clean[-1]
                if last_char in "。;；?？!！:：":
                    ends_with_sentence_terminator = True
                elif last_char == '.':
                    if len(prev_clean) >= 2 and prev_clean[-2].isdigit():
                        ends_with_sentence_terminator = False
                    else:
                        ends_with_sentence_terminator = True
            
            # 如果当前行不是一个新的 bullet/序列号，且上一行没有句末终结标点，则合并
            if not has_bullet_prefix and not has_numeric_prefix and not ends_with_sentence_terminator:
                if prev_line and prev_line[-1].isalnum() and line and line[0].isalnum():
                    merged_job_lines[-1] = prev_line + " " + line
                else:
                    merged_job_lines[-1] = prev_line + line
            else:
                merged_job_lines.append(line)

        # 3. 对合并后的完整行进行清洗与过滤，生成最终的 bullets
        bullets: list[str] = []
        for raw in merged_job_lines:
            text = _BULLET_PREFIX_RE.sub("", raw).strip()
            if not text:
                continue
            has_bullet_prefix = bool(_BULLET_PREFIX_RE.match(raw))
            # 启发式过滤：这些行不是 bullet
            # 1) 页脚个人信息：含手机号或 @ 邮箱
            digit_count = sum(1 for c in text if c.isdigit())
            if digit_count >= 8 or "@" in text:
                continue
            # 2) 极短纯中文 (< 4 字)：不可能是 bullet，常是姓名/小标题
            if len(text) <= 4 and all(0x4e00 <= ord(c) <= 0x9fff for c in text):
                continue
            # 3) 堆栈清单：含 "："/":" 起头 + 多个分隔符
            if (text.startswith("关键技术") or text.startswith("技术栈") or text.startswith("Skills")):
                continue
            # 4) 段落 sub-header 短行：不含任何句末/句中标点的行更像小标题而非描述
            #    例: "能力部署及大模型推理平台任务部署方式升级" / "GPU资源优化与切片技术落地"
            #    例外：有 • 前缀的裸行（如 "• 主导稳定性建设" 无标点也算 bullet）
            if not has_bullet_prefix:
                sentence_punct = set("。，；：、？！（）()")
                if not any(c in sentence_punct for c in text):
                    continue
            bullets.append(text)

        jobs.append({
            "company": company or "",
            "role": role or "",
            "period": period,
            "bullets": bullets,
        })

    return jobs


_EDU_SECTION_KEYS = ("教育背景", "教育经历", "教育信息", "education", "academic background")
_PROJ_SECTION_KEYS = ("项目经历", "项目经验", "项目案例", "个人项目", "projects", "personal projects", "project experience")
_SKILL_SECTION_KEYS = ("专业技能", "技能", "技能特长", "IT技能", "skills", "technical skills", "professional skills")
_DEGREE_KEYWORDS = ("学士", "硕士", "博士", "本科", "大专", "专科", "研究生", "高中", "master", "bachelor", "phd", "b.s.", "m.s.", "b.a.", "m.a.")


def _parse_education(lines: list[str]) -> list[dict]:
    if not lines:
        return []
    header_line_idx = []
    header_date = []
    for i, line in enumerate(lines):
        clean = line.strip()
        if not clean:
            continue
        m = _DATE_RE.search(clean)
        if m or any(k in clean for k in ("大学", "学院", "分校", "University", "College", "Institute", "School")):
            header_line_idx.append(i)
            header_date.append(m)

    if not header_line_idx:
        bullets = [l.strip() for l in lines if l.strip() and not any(_clean_line(l.strip()).lower().startswith(k) for k in _NEXT_SECTION_KEYS)]
        if bullets:
            return [{"school": bullets[0], "degree": "", "major": "", "period": "", "bullets": bullets[1:]}]
        return []

    entries = []
    for n, idx in enumerate(header_line_idx):
        line = lines[idx].strip()
        date_m = header_date[n]
        period = date_m.group(0).strip() if date_m else ""
        before = line[:date_m.start()].strip(" \t|-—–·") if date_m else line
        parts = [p.strip() for p in re.split(r"[\s　]{2,}|\t+|[|｜]", before) if p.strip()]

        school, degree, major = "", "", ""
        for p in parts:
            p_lower = p.lower()
            if any(d in p_lower for d in _DEGREE_KEYWORDS):
                degree = p
            elif any(u in p for u in ("大学", "学院", "分校", "University", "College", "Institute", "School")):
                school = p
            elif not school:
                school = p
            elif not major:
                major = p

        end_idx = header_line_idx[n + 1] if n + 1 < len(header_line_idx) else len(lines)
        bullets = []
        for j in range(idx + 1, end_idx):
            raw = lines[j].strip()
            if not raw or any(_clean_line(raw).lower().startswith(k) for k in _NEXT_SECTION_KEYS):
                continue
            bullets.append(_BULLET_PREFIX_RE.sub("", raw).strip())

        entries.append({
            "school": school,
            "degree": degree,
            "major": major,
            "period": period,
            "bullets": bullets,
        })
    return entries


def _parse_projects(lines: list[str]) -> list[dict]:
    if not lines:
        return []
    header_line_idx = []
    header_date = []
    for i, line in enumerate(lines):
        clean = line.strip()
        if not clean:
            continue
        m = _DATE_RE.search(clean)
        if m or clean.startswith("【") or clean.startswith("[") or "项目" in clean or "System" in clean or "Platform" in clean:
            if not _BULLET_PREFIX_RE.match(clean) and len(clean) < 80:
                header_line_idx.append(i)
                header_date.append(m)

    if not header_line_idx:
        bullets = [_BULLET_PREFIX_RE.sub("", l.strip()).strip() for l in lines if l.strip() and not _is_header_match(l, _NEXT_SECTION_KEYS)]
        if bullets:
            return [{
                "name": bullets[0],
                "role": "",
                "period": "",
                "tech_stack": "",
                "bullets": bullets[1:] if len(bullets) > 1 else bullets,
            }]
        return []

    projects = []
    for n, idx in enumerate(header_line_idx):
        line = lines[idx].strip()
        date_m = header_date[n]
        period = date_m.group(0).strip() if date_m else ""
        before = line[:date_m.start()].strip(" \t|-—–·") if date_m else line

        parts = [p.strip() for p in re.split(r"[\s　]{2,}|\t+|[|｜]", before) if p.strip()]
        name = parts[0] if parts else before
        role = parts[1] if len(parts) > 1 else ""

        end_idx = header_line_idx[n + 1] if n + 1 < len(header_line_idx) else len(lines)
        bullets = []
        tech_stack = ""
        for j in range(idx + 1, end_idx):
            raw = lines[j].strip()
            if not raw or any(_clean_line(raw).lower().startswith(k) for k in _NEXT_SECTION_KEYS):
                continue
            clean_b = _BULLET_PREFIX_RE.sub("", raw).strip()
            if clean_b.startswith("技术栈") or clean_b.startswith("关键技术") or clean_b.startswith("Tech Stack"):
                tech_stack = re.sub(r"^(?:技术栈|关键技术|Tech Stack)[:：\s]*", "", clean_b)
            else:
                bullets.append(clean_b)

        projects.append({
            "name": name,
            "role": role,
            "period": period,
            "tech_stack": tech_stack,
            "bullets": bullets,
        })
    return projects


def _parse_skills(lines: list[str]) -> list[str]:
    skills = []
    for line in lines:
        raw = line.strip()
        if not raw or any(_clean_line(raw).lower().startswith(k) for k in _NEXT_SECTION_KEYS):
            continue
        clean = _BULLET_PREFIX_RE.sub("", raw).strip()
        if clean:
            if "：" in clean or ":" in clean or len(clean) > 25:
                if clean not in skills:
                    skills.append(clean)
            else:
                items = [item.strip() for item in re.split(r"[,;；\n•·|]+", clean) if item.strip()]
                for it in items:
                    if len(it) < 60 and it not in skills:
                        skills.append(it)
    return skills


_SUMMARY_SECTION_KEYS = ("自我评价", "个人评价", "个人总结", "综合评价", "summary", "personal summary")


def _parse_summary(lines: list[str]) -> str:
    collected = []
    for line in lines:
        raw = line.strip()
        if not raw or any(_clean_line(raw).lower().startswith(k) for k in _NEXT_SECTION_KEYS):
            continue
        collected.append(_BULLET_PREFIX_RE.sub("", raw).strip())
    return "\n".join(collected)


def parse_resume_structure(text: str) -> dict:
    """主入口。返回 {profile, work_experiences, education, projects, skills, summary}，所有字段原文 verbatim。"""
    lines = text.split("\n")

    work_start, work_end = _find_section_bounds(lines, _WORK_SECTION_KEYS)
    intern_start, intern_end = _find_section_bounds(lines, _INTERN_SECTION_KEYS)
    edu_start, edu_end = _find_section_bounds(lines, _EDU_SECTION_KEYS)
    proj_start, proj_end = _find_section_bounds(lines, _PROJ_SECTION_KEYS)
    skill_start, skill_end = _find_section_bounds(lines, _SKILL_SECTION_KEYS)
    sum_start, sum_end = _find_section_bounds(lines, _SUMMARY_SECTION_KEYS)

    work_experiences: list[dict] = []
    if work_start is not None:
        work_section = lines[work_start + 1: work_end]
        work_experiences.extend(_parse_work_experiences(work_section))
    if intern_start is not None:
        if work_start is None or intern_start < work_start or intern_start >= (work_end or 0):
            intern_section = lines[intern_start + 1: intern_end]
            work_experiences.extend(_parse_work_experiences(intern_section))

    education = _parse_education(lines[edu_start + 1: edu_end]) if edu_start is not None else []
    projects = _parse_projects(lines[proj_start + 1: proj_end]) if proj_start is not None else []
    skills = _parse_skills(lines[skill_start + 1: skill_end]) if skill_start is not None else []
    summary = _parse_summary(lines[sum_start + 1: sum_end]) if sum_start is not None else ""
    profile = _parse_profile(lines, text)

    return {
        "profile": profile,
        "work_experiences": work_experiences,
        "education": education,
        "projects": projects,
        "skills": skills,
        "summary": summary,
    }

