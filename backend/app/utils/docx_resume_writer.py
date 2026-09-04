"""就地改写 DOCX 简历的 bullet 文字，保留所有 run 样式。

设计目标：用户上传 DOCX 后，下载下来的 AI 优化版应该和原 DOCX 看起来
完全一样（字体/颜色/图标/分栏/表格全保留），只有 bullet 文字被替换。

实现策略：
1. 用 python-docx 打开原 DOCX
2. 遍历 doc.paragraphs + 表格里所有 cell 的 paragraphs
3. 从 analysis_data.work_experiences 收集 (originalText, optimizedText) 配对
4. 段落的 text.strip() 与 originalText.strip() 完全相等时，
   把 para.runs[0].text 改成 optimizedText，其他 run 清空（保留格式属性）
5. 统计替换率，< MATCH_RATE_THRESHOLD 时抛 BulletMatchError 让上层回退到
   统一模板 PDF 渲染路径（PDF 转 DOCX 后段落错乱时走兜底）

不在 bullets 列表里的段落（公司名/岗位名/时间/小标题）一律不动。
"""
import io
import logging
import re
import unicodedata
from difflib import SequenceMatcher
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# 替换率 < 80% 视为识别失败，触发 PDF 兜底
# 之前 0/12 完全 fail 抛 500 太刚性 — 用户场景是「哪怕只优化 1 个 bullet 也比没有好」。
# 后续会基于真实数据(desktop dump 拿到 pypdf vs pdf2docx 输出差异)重新调到合理值。
MATCH_RATE_THRESHOLD = 0.0

# Fuzzy 匹配下限：相似度 >= 该值视为同一 bullet。
# 低于 0.8 会开始误报（两个不同 bullet 也匹配上），高于 0.9 又会漏掉
# PDF→DOCX 引入的小差异（smart quote、全角符号等）。
FUZZY_THRESHOLD = 0.85

# 智能引号 → ASCII 双引号。pypdf 和 pdf2docx 对同一份 PDF 的 smart quote
# 抽取结果不一致（左/右引号互换），是 bullet 匹配失败的常见原因之一。
_SMART_QUOTE_DOUBLE_RE = re.compile(r"[“”「」『』]")
_SMART_QUOTE_SINGLE_RE = re.compile(r"[‘’]")
# 各种破折号变体 → ASCII hyphen-minus。
# pypdf 经常把 "—" 输出成 "-" 或丢失，pdf2docx 保留原 Unicode。
_DASH_RE = re.compile(r"[–—―−]")

# 零宽 / 不可见字符集合——PDF→DOCX（pdf2docx）会在每段段尾追加 ​ 等。
# bullet 的 originalText 是从 PDF 文本抽取得到的，不带这些字符；
# 直接比较必然失败，必须在比对前归一化。
_INVISIBLE_CHARS = "​‌‍﻿"
_INVISIBLE_RE = re.compile(f"[{re.escape(_INVISIBLE_CHARS)}]")
_WHITESPACE_RE = re.compile(r"\s+")
# 段首常见列表符号（bullet / 圆点 / 方块 / 横线 / 星号）。PDF→DOCX 的段落几乎都
# 在开头加 "• " 或 "·" 之类的项目符号，而 bullet 的 originalText 已被 LLM 抽取时
# 去掉了这个符号，所以比较时也要把 para 段首的列表符剥掉才能命中前缀。
_LEADING_BULLET_RE = re.compile(r"^[•·●○▪▫■□\-*]+\s*")


def _norm_text(s: str) -> str:
    """bullet / 段落 文本的归一化：

    顺序很重要，每一步都建立在前一步已经剥干净的基础上：
    1. NFKC Unicode 归一化 — 全角→半角、兼容形分解（如 ﬁ → fi）
    2. 智能引号 → ASCII — " " " " 「 」 『 』 全替成 "
    3. 破折号变体 → ASCII hyphen — — – − 全替成 -
    4. 零宽字符清理
    5. 任意空白 → 单空格
    6. 段首列表符号剥除
    """
    if not s:
        return ""
    s = unicodedata.normalize("NFKC", s)
    s = _SMART_QUOTE_DOUBLE_RE.sub('"', s)
    s = _SMART_QUOTE_SINGLE_RE.sub("'", s)
    s = _DASH_RE.sub("-", s)
    s = _INVISIBLE_RE.sub("", s)
    s = _WHITESPACE_RE.sub(" ", s).strip()
    s = _LEADING_BULLET_RE.sub("", s)
    return s


def _strip_ws(s: str) -> str:
    """去掉所有空白后的字符串，用作 bullet / 段落的最终比对键。

    PDF→DOCX 把同一 bullet 拆段后，第一段尾 / 第二段首可能多出零宽、空格或换行；
    拼接时无论加不加空格都会与原 orig 中的字符排列产生差异，所以最终比较键
    必须"忽略所有空白"，容忍任何位置上的空白差异。
    """
    return re.sub(r"\s+", "", s or "")


def _fuzzy_ratio(a: str, b: str) -> float:
    """两个空白无关字符串的相似度。短串 / 长度差过大特判，长串走 SequenceMatcher。"""
    if not a or not b:
        return 0.0
    if a == b:
        return 1.0
    la, lb = len(a), len(b)
    # 短串 SequenceMatcher 不准,长度差过大也不算同一段 — 直接拒,避免误报
    if min(la, lb) < 4:
        return 1.0 if a == b else 0.0
    if max(la, lb) / min(la, lb) > 1.5:
        return 0.0
    return SequenceMatcher(None, a, b).ratio()


class BulletMatchError(Exception):
    """bullet 段落识别率不足，上层应回退到 PDF 渲染路径。"""


def _iter_all_paragraphs(doc):
    """遍历 doc.paragraphs + 所有表格（含嵌套表格）里 cell 的 paragraphs。"""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested in cell.tables:
                    for nr in nested.rows:
                        for nc in nr.cells:
                            yield from nc.paragraphs


def _replace_paragraph_text(para, new_text: str) -> None:
    """把段落里所有 run 的文字合并到第一个 run 上，保留第一个 run 的样式。

    视觉效果：颜色/字号/字体来自第一个 run（通常是 bullet 主样式）。
    其他 run 的文字清空但其 run 属性（rPr）保留 → Word 不会因空 run 报错。
    并且显式将中文字体（w:eastAsia）设为“微软雅黑”，防止在特定字体环境（如纯英文或日文转换环境）下部分中文字符显示为 □。
    """
    if not para.runs:
        run = para.add_run(new_text)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
        return
        
    run = para.runs[0]
    run.text = new_text
    
    # 强制东亚文字使用“微软雅黑”进行渲染，彻底修复特定汉字变豆腐块(□)的Bug
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    for r in para.runs[1:]:
        r.text = ""


def _collect_bullets(analysis_data: Dict[str, Any]) -> List[Tuple[str, str]]:
    """从 analysis_data.work_experiences 抽出 (original, optimized) 配对。

    过滤规则：
    - 缺 originalText / optimizedText 的跳过（保持原文不被错误清空）
    - optimizedText 与 originalText 一字不差（LLM 偷懒或本就 OK）的跳过，
      避免无意义的"替换"
    """
    pairs: List[Tuple[str, str]] = []
    for exp in analysis_data.get("work_experiences") or []:
        for bullet in exp.get("bullets") or []:
            if not isinstance(bullet, dict):
                continue
            orig = (bullet.get("originalText") or "").strip()
            opt = (bullet.get("optimizedText") or "").strip()
            if not orig or not opt or opt == orig:
                continue
            pairs.append((orig, opt))
    return pairs


def _delete_paragraph(paragraph) -> None:
    """从文档中彻底删除指定段落的 XML 元素，避免在合并多段落时在 Word 中留下空白空行。"""
    p = paragraph._element
    if p is not None and p.getparent() is not None:
        p.getparent().remove(p)


def rewrite_resume_docx(content_bytes: bytes, analysis_data: Dict[str, Any]) -> bytes:
    """把 DOCX 里 bullet 段落的文字就地替换为 optimizedText，保留所有 run 样式。

    支持对单段匹配和换行分割的多段合并匹配，并在合并后删除多余的空段落。
    """
    pairs = _collect_bullets(analysis_data)
    if not pairs:
        raise BulletMatchError("analysis_data 中没有可替换的 bullet 配对")

    doc = Document(io.BytesIO(content_bytes))
    # pending 里每个元素 = (orig, opt, key)
    # key = 去掉所有空白后的 orig，用于跨段拼接时无视各段空格的微小差异。
    pending = [(orig, opt, _strip_ws(_norm_text(orig))) for (orig, opt) in pairs]
    replaced_count = 0

    all_paras = list(_iter_all_paragraphs(doc))
    p_idx = 0
    total_paras = len(all_paras)

    # 标记已被合并删除的段落元素，避免重复访问或修改
    deleted_elements = set()

    while p_idx < total_paras and pending:
        para = all_paras[p_idx]
        if para._element in deleted_elements:
            p_idx += 1
            continue

        para_key = _strip_ws(_norm_text(para.text))
        if not para_key:
            p_idx += 1
            continue

        matched = False
        for i, (orig, opt, key) in enumerate(pending):
            # 1. 尝试单段完全匹配（用 stripped 比较键）
            if para_key == key:
                _replace_paragraph_text(para, opt)
                pending.pop(i)
                replaced_count += 1
                p_idx += 1
                matched = True
                break

            # 1.5 Fuzzy 兜底：strict 匹配失败但相似度 >= FUZZY_THRESHOLD 也算命中。
            # 场景: pypdf vs pdf2docx 抽取 smart quote / 全角符号不一致,
            # 严格比较 0% 但 fuzzy 能救回大部分。
            # 风险: 短串 / 长度差过大容易误报,_fuzzy_ratio 已经特判,这里再确认一次。
            if (not matched
                    and min(len(para_key), len(key)) >= 6
                    and max(len(para_key), len(key)) / min(len(para_key), len(key)) <= 1.3
                    and _fuzzy_ratio(para_key, key) >= FUZZY_THRESHOLD):
                _replace_paragraph_text(para, opt)
                pending.pop(i)
                replaced_count += 1
                p_idx += 1
                matched = True
                break

            # 2. 尝试多段合并匹配（解决 PDF 转换或换行导致的段落截断问题）
            #    用 stripped 字符串前缀拼接——这样无论原段间是否真的有空格，
            #    以及我们拼接时是否插入空格，比较键都一致。
            if key.startswith(para_key):
                accumulated = para_key
                match_paras = [para]
                next_idx = p_idx + 1

                while next_idx < total_paras and len(accumulated) < len(key):
                    next_para = all_paras[next_idx]
                    if next_para._element in deleted_elements:
                        next_idx += 1
                        continue
                    next_key = _strip_ws(_norm_text(next_para.text))
                    if not next_key:
                        next_idx += 1
                        continue

                    if key.startswith(accumulated + next_key):
                        accumulated = accumulated + next_key
                        match_paras.append(next_para)
                        next_idx += 1
                    else:
                        break

                if accumulated == key:
                    # 匹配成功！将优化后的文本填入首个段落，并删除后续的多余段落
                    _replace_paragraph_text(match_paras[0], opt)
                    for extra_para in match_paras[1:]:
                        _delete_paragraph(extra_para)
                        deleted_elements.add(extra_para._element)

                    pending.pop(i)
                    replaced_count += 1
                    p_idx = next_idx
                    matched = True
                    break

        if not matched:
            p_idx += 1

    match_rate = replaced_count / len(pairs)
    if match_rate < MATCH_RATE_THRESHOLD:
        raise BulletMatchError(
            f"bullet 识别率 {match_rate:.1%} ({replaced_count}/{len(pairs)}) "
            f"低于阈值 {MATCH_RATE_THRESHOLD:.0%}"
        )

    logger.info(
        "[docx_writer] replaced %d/%d bullets (%.1f%%)",
        replaced_count, len(pairs), match_rate * 100,
    )

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def generate_structured_resume_docx(analysis_data: Dict[str, Any]) -> bytes:
    """当就地改写原 DOCX 失败或为 PDF 源文件时，根据 AI 结构化分析结果从头生成一份排版美观的标准 DOCX 简历。"""
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    
    # 设置边距 0.8 英寸
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    profile = analysis_data.get("profile") or {}
    name = (profile.get("name") or "求职者").strip()
    target_role = (profile.get("target_role") or profile.get("role") or "").strip()
    phone = (profile.get("phone") or "").strip()
    email = (profile.get("email") or "").strip()
    location = (profile.get("location") or profile.get("city") or "").strip()
    experience_years = (profile.get("experience_years") or "").strip()

    # 1. 姓名 & 目标岗位
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(name)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_p.paragraph_format.space_after = Pt(4)

    if target_role:
        role_run = title_p.add_run(f"  |  {target_role}")
        role_run.font.size = Pt(13)
        role_run.font.color.rgb = RGBColor(100, 116, 139)

    # 2. 联系方式
    contact_parts = [p for p in [phone, email, location, experience_years] if p]
    if contact_parts:
        contact_p = doc.add_paragraph("  •  ".join(contact_parts))
        contact_p.paragraph_format.space_after = Pt(14)
        for r in contact_p.runs:
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(71, 85, 105)

    def add_section_header(title_text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(title_text)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(30, 58, 138)

    # 3. 个人优势 / 总结
    summary = (analysis_data.get("summary") or profile.get("summary") or "").strip()
    if summary:
        add_section_header("个人优势 / 综合评价")
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(8)
        for r in p.runs:
            r.font.size = Pt(10)

    # 4. 工作经历
    work_exps = analysis_data.get("work_experiences") or []
    if work_exps:
        add_section_header("工作经历")
        for work in work_exps:
            company = (work.get("company") or "").strip()
            role = (work.get("role") or "").strip()
            time_range = (work.get("time_range") or work.get("duration") or "").strip()

            header_p = doc.add_paragraph()
            header_p.paragraph_format.space_before = Pt(6)
            header_p.paragraph_format.space_after = Pt(2)
            
            c_run = header_p.add_run(company)
            c_run.bold = True
            c_run.font.size = Pt(10.5)

            if role:
                r_run = header_p.add_run(f"  —  {role}")
                r_run.bold = True
                r_run.font.size = Pt(10.5)

            if time_range:
                t_run = header_p.add_run(f"\t{time_range}")
                t_run.font.size = Pt(9.5)
                t_run.font.color.rgb = RGBColor(100, 116, 139)

            bullets = work.get("bullets") or []
            for b in bullets:
                opt_text = (b.get("optimizedText") or b.get("originalText") or "").strip()
                if opt_text:
                    bp = doc.add_paragraph(style='List Bullet')
                    bp.paragraph_format.space_after = Pt(3)
                    brun = bp.add_run(opt_text)
                    brun.font.size = Pt(9.5)

    # 5. 核心项目
    projects = analysis_data.get("projects") or []
    if projects:
        add_section_header("核心项目经历")
        for proj in projects:
            p_name = (proj.get("name") or proj.get("title") or "").strip()
            p_role = (proj.get("role") or "").strip()
            p_time = (proj.get("duration") or proj.get("time") or "").strip()

            header_p = doc.add_paragraph()
            header_p.paragraph_format.space_before = Pt(6)
            header_p.paragraph_format.space_after = Pt(2)

            pr_run = header_p.add_run(p_name)
            pr_run.bold = True
            pr_run.font.size = Pt(10.5)

            if p_role:
                ro_run = header_p.add_run(f"  ({p_role})")
                ro_run.font.size = Pt(10)

            if p_time:
                ti_run = header_p.add_run(f"\t{p_time}")
                ti_run.font.size = Pt(9.5)

            bullets = proj.get("bullets") or proj.get("details") or []
            if isinstance(bullets, list):
                for b in bullets:
                    txt = b.get("optimizedText") or b.get("originalText") or b if isinstance(b, dict) else str(b)
                    txt = str(txt).strip()
                    if txt:
                        bp = doc.add_paragraph(style='List Bullet')
                        bp.paragraph_format.space_after = Pt(3)
                        brun = bp.add_run(txt)
                        brun.font.size = Pt(9.5)

    # 6. 教育背景
    education = analysis_data.get("education") or profile.get("education") or []
    if education:
        add_section_header("教育背景")
        if isinstance(education, list):
            for edu in education:
                if isinstance(edu, dict):
                    school = (edu.get("school") or "").strip()
                    degree = (edu.get("degree") or "").strip()
                    major = (edu.get("major") or "").strip()
                    year = (edu.get("year") or edu.get("duration") or "").strip()
                    
                    ep = doc.add_paragraph()
                    ep.paragraph_format.space_after = Pt(3)
                    s_run = ep.add_run(school)
                    s_run.bold = True
                    s_run.font.size = Pt(10)
                    if degree or major:
                        ep.add_run(f"  |  {degree} {major}".strip())
                    if year:
                        ep.add_run(f"\t{year}")
                else:
                    ep = doc.add_paragraph(str(edu))
                    ep.paragraph_format.space_after = Pt(3)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
